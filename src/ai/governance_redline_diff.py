"""
Governance Redline Diff Engine (CHUNK 6.21.1 - Part 15)
Compares two governance documents and generates redline/track-changes style documents.
"""
import difflib
from typing import Dict, List, Any, Optional
from pathlib import Path

try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from .medical_llm import call_medical_llm
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False


class GovernanceRedlineDiffEngine:
    """
    Compares two governance packages (v1 vs v2) and generates:
    - Redline text diff
    - Structured list of added/removed elements
    - AI-generated interpretation of changes
    - Editable DOCX track-changes-like document
    """

    def __init__(self):
        """Initialize the Redline Diff Engine."""
        pass

    def generate_diff(self, old_text: str, new_text: str) -> Dict[str, Any]:
        """
        Generate a structured diff between two text documents.
        
        Args:
            old_text: Original version text
            new_text: New version text
            
        Returns:
            Dictionary with added, removed, and unchanged sections
        """
        if not old_text:
            old_text = ""
        if not new_text:
            new_text = ""
        
        diff = difflib.ndiff(old_text.splitlines(keepends=True), new_text.splitlines(keepends=True))
        
        added = []
        removed = []
        unchanged = []
        
        for line in diff:
            if line.startswith("+ "):
                added.append(line[2:].rstrip('\n'))
            elif line.startswith("- "):
                removed.append(line[2:].rstrip('\n'))
            elif line.startswith("  "):
                unchanged.append(line[2:].rstrip('\n'))
        
        return {
            "added": added,
            "removed": removed,
            "unchanged": unchanged,
            "total_added": len(added),
            "total_removed": len(removed),
            "net_change": len(added) - len(removed)
        }

    def generate_redline_docx(self, diff_dict: Dict[str, Any], output_path: str) -> str:
        """
        Generate a DOCX document with redline/track-changes formatting.
        
        Args:
            diff_dict: Dictionary from generate_diff()
            output_path: Output file path
            
        Returns:
            Path to generated DOCX
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install it with: pip install python-docx"
            )
        
        doc = Document()
        
        # Title
        title = doc.add_heading("Governance Redline Comparison", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("Auto-generated change log comparing governance document versions.")
        doc.add_paragraph(f"Total additions: {diff_dict.get('total_added', 0)}")
        doc.add_paragraph(f"Total removals: {diff_dict.get('total_removed', 0)}")
        doc.add_paragraph(f"Net change: {diff_dict.get('net_change', 0)} lines")
        doc.add_paragraph("")
        
        # Added text → green
        if diff_dict.get("added"):
            doc.add_heading("Additions (New Content)", level=1)
            for line in diff_dict["added"]:
                if line.strip():  # Skip empty lines
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.color.rgb = RGBColor(0, 150, 0)  # green
            doc.add_paragraph("")
        
        # Removed text → red + strikethrough
        if diff_dict.get("removed"):
            doc.add_heading("Removals (Deleted Content)", level=1)
            for line in diff_dict["removed"]:
                if line.strip():  # Skip empty lines
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.color.rgb = RGBColor(180, 0, 0)  # red
                    run.font.strike = True
            doc.add_paragraph("")
        
        # Unchanged context (optional, can be large)
        if diff_dict.get("unchanged") and len(diff_dict["unchanged"]) < 50:
            doc.add_heading("Unchanged Content (Context)", level=1)
            for line in diff_dict["unchanged"][:20]:  # Limit context
                if line.strip():
                    doc.add_paragraph(line)
        
        doc.save(output_path)
        return output_path

    def generate_ai_interpretation(self, old_text: str, new_text: str, 
                                   diff_dict: Optional[Dict[str, Any]] = None) -> str:
        """
        Generate AI-powered interpretation of changes.
        
        Args:
            old_text: Original version text
            new_text: New version text
            diff_dict: Optional pre-computed diff dictionary
            
        Returns:
            AI-generated interpretation text
        """
        if not LLM_AVAILABLE:
            return "AI interpretation unavailable. Install medical_llm module for LLM-powered change analysis."
        
        if diff_dict is None:
            diff_dict = self.generate_diff(old_text, new_text)
        
        # Create summary of changes
        change_summary = f"""
Summary of Changes:
- {diff_dict.get('total_added', 0)} lines added
- {diff_dict.get('total_removed', 0)} lines removed
- Net change: {diff_dict.get('net_change', 0)} lines

Key Additions:
{chr(10).join(diff_dict.get('added', [])[:10])}

Key Removals:
{chr(10).join(diff_dict.get('removed', [])[:10])}
"""
        
        prompt = f"""
You are a pharmacovigilance inspector reviewing governance document changes.

Compare two versions of a signal governance document.

OLD VERSION:
{old_text[:2000]}

NEW VERSION:
{new_text[:2000]}

CHANGE SUMMARY:
{change_summary}

Explain:
1. What changed between the two versions?
2. Why these changes matter for signal governance?
3. Impact on medical judgment and risk assessment
4. Impact on risk management decisions
5. Whether the changes are aligned with GVP Module IX and FDA expectations
6. Any regulatory concerns or improvements needed

Provide a professional, regulatory-compliant analysis suitable for inspection documentation.
"""
        
        try:
            system_prompt = "You are a pharmacovigilance inspector reviewing governance document changes for regulatory compliance."
            return call_medical_llm(
                prompt=prompt,
                system_prompt=system_prompt,
                task_type="general",
                max_tokens=1500,
                temperature=0.3
            ) or "AI interpretation unavailable."
        except Exception as e:
            return f"AI interpretation error: {str(e)}"

    def run(self, old_text: str, new_text: str, output_path: str = "./redline.docx") -> Dict[str, Any]:
        """
        Complete redline comparison workflow.
        
        Args:
            old_text: Original version text
            new_text: New version text
            output_path: Output DOCX file path
            
        Returns:
            Dictionary with diff, docx path, and AI summary
        """
        # Generate diff
        diff = self.generate_diff(old_text, new_text)
        
        # Generate DOCX redline
        docx_path = None
        if DOCX_AVAILABLE:
            try:
                docx_path = self.generate_redline_docx(diff, output_path)
            except Exception as e:
                docx_path = f"Error generating DOCX: {str(e)}"
        
        # Generate AI interpretation
        ai_summary = self.generate_ai_interpretation(old_text, new_text, diff)
        
        return {
            "diff": diff,
            "docx": docx_path,
            "summary": ai_summary,
            "total_added": diff.get("total_added", 0),
            "total_removed": diff.get("total_removed", 0),
            "net_change": diff.get("net_change", 0)
        }

    def compare_signal_files(self, old_signal_file: Dict[str, Any], 
                            new_signal_file: Dict[str, Any]) -> Dict[str, Any]:
        """
        Compare two signal file dictionaries and generate redline.
        
        Args:
            old_signal_file: Old signal file dictionary
            new_signal_file: New signal file dictionary
            
        Returns:
            Comparison dictionary with section-by-section diffs
        """
        sections = ["overview", "evidence_package", "ai_summary", "followup_plan"]
        section_diffs = {}
        
        for section in sections:
            old_section = str(old_signal_file.get(section, ""))
            new_section = str(new_signal_file.get(section, ""))
            
            if old_section or new_section:
                section_diffs[section] = self.generate_diff(old_section, new_section)
        
        return {
            "section_diffs": section_diffs,
            "sections_compared": list(section_diffs.keys())
        }

