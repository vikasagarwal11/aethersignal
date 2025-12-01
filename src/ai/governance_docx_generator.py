"""
Governance DOCX Generator (CHUNK 6.21.1 - Part 14)
Generates regulatory-ready DOCX documents for signal governance files.
"""
import datetime
from typing import Dict, Any, Optional
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .signal_file_generator import SignalFileGenerator


class GovernanceDOCXGenerator:
    """
    Generates a fully editable DOCX governance package including:
    - Signal Overview
    - Trend Alerts
    - Compliance Checklist
    - Risk Profile (RPF)
    - Reviewer Assignment
    - Follow-Up Plan
    - AI Summary
    - Governance Notes
    """

    def __init__(self, dataset: Dict[str, Any]):
        """
        Initialize the DOCX Generator.
        
        Args:
            dataset: Dictionary containing signals and data
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install it with: pip install python-docx"
            )
        
        self.dataset = dataset
        self.signal_file = SignalFileGenerator(dataset)

    def generate_docx(self, signal_id: str, output_path: str) -> str:
        """
        Generate complete DOCX for a signal.
        
        Args:
            signal_id: Signal identifier
            output_path: Output file path
            
        Returns:
            Path to generated DOCX
        """
        doc = Document()
        
        # Generate signal file data
        signal_data = self.signal_file.generate_signal_file(signal_id)
        
        if "error" in signal_data:
            raise ValueError(f"Signal not found: {signal_data.get('error')}")
        
        # Title
        title = doc.add_heading("Signal Governance Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated at: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
        doc.add_paragraph("")
        
        # Section 1: Signal Overview
        doc.add_heading("1. Signal Overview", level=1)
        overview = signal_data.get("overview", {})
        evidence = signal_data.get("evidence_package", {})
        
        overview_items = [
            ("Drug", evidence.get("drug", overview.get("drug", "Unknown"))),
            ("Reaction", evidence.get("reaction", overview.get("reaction", "Unknown"))),
            ("Cases", str(evidence.get("cases", overview.get("cases", 0)))),
            ("Serious Cases", str(evidence.get("serious_cases", overview.get("serious_cases", 0)))),
            ("Fatal Cases", str(evidence.get("fatal_cases", overview.get("fatal_cases", 0)))),
            ("Lifecycle", evidence.get("lifecycle", overview.get("lifecycle", "Unknown"))),
            ("Detected On", evidence.get("detected_on", overview.get("detected_on", "Unknown"))),
        ]
        
        for key, value in overview_items:
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(str(value))
        
        doc.add_paragraph("")
        
        # Section 2: Trend Alerts Summary
        doc.add_heading("2. Trend Alerts Summary", level=1)
        trends = signal_data.get("trends", {})
        
        alerts = trends.get("alerts", [])
        if alerts:
            for alert in alerts[:10]:  # Limit to first 10
                if isinstance(alert, dict):
                    title = alert.get("title", alert.get("summary", "Alert"))
                    severity = alert.get("severity", "info")
                    summary = alert.get("summary", "")
                    p = doc.add_paragraph()
                    p.add_run(f"{severity.upper()}: {title}").bold = True
                    doc.add_paragraph(summary)
        else:
            doc.add_paragraph("No significant alerts detected.")
        
        doc.add_paragraph("")
        
        # Section 3: Compliance Checklist
        doc.add_heading("3. Compliance Checklist (GVP / FDA / MHRA)", level=1)
        compliance = signal_data.get("compliance", {})
        checklist = compliance.get("checklist", [])
        
        if checklist:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Requirement"
            hdr_cells[1].text = "Category"
            hdr_cells[2].text = "Regulation"
            hdr_cells[3].text = "Status"
            
            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add data rows
            for item in checklist:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get("text", "")
                row_cells[1].text = item.get("category", "")
                row_cells[2].text = item.get("regulation", "")
                completed = item.get("completed", False)
                row_cells[3].text = "✓ Passed" if completed else "✗ Gap"
            
            doc.add_paragraph("")
            
            # Compliance score
            score = compliance.get("score", 0)
            p = doc.add_paragraph()
            p.add_run(f"Overall Compliance Score: ").bold = True
            p.add_run(f"{score:.1f}/100")
        else:
            doc.add_paragraph("No compliance checklist available.")
        
        doc.add_paragraph("")
        
        # Section 4: Risk Profile
        doc.add_heading("4. Risk Assessment (RPF)", level=1)
        risk_profile = signal_data.get("risk_profile", {})
        
        p = doc.add_paragraph()
        p.add_run("RPF Score: ").bold = True
        p.add_run(str(risk_profile.get("rpf_score", 0)))
        
        p = doc.add_paragraph()
        p.add_run("Risk Level: ").bold = True
        p.add_run(str(risk_profile.get("risk_level", "Low")))
        
        sub_scores = risk_profile.get("sub_scores", {})
        if sub_scores:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.add_run("Sub-Scores:").bold = True
            for key, value in sub_scores.items():
                doc.add_paragraph(
                    f"  • {key.replace('_', ' ').title()}: {value:.2f}",
                    style='List Bullet'
                )
        
        doc.add_paragraph("")
        
        # Section 5: Reviewer Assignment
        doc.add_heading("5. Reviewer Assignment", level=1)
        reviewer = signal_data.get("reviewer_assignment", {})
        
        p = doc.add_paragraph()
        p.add_run("Assigned Reviewer: ").bold = True
        p.add_run(reviewer.get("reviewer", "Unassigned"))
        
        reasons = reviewer.get("reason", [])
        if isinstance(reasons, list) and reasons:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.add_run("Assignment Reasons:").bold = True
            for reason in reasons:
                doc.add_paragraph(f"  • {reason}", style='List Bullet')
        elif isinstance(reasons, str):
            p = doc.add_paragraph()
            p.add_run("Reason: ").bold = True
            p.add_run(reasons)
        
        workload_status = reviewer.get("workload_status", "unknown")
        p = doc.add_paragraph()
        p.add_run("Workload Status: ").bold = True
        p.add_run(workload_status)
        
        doc.add_paragraph("")
        
        # Section 6: Follow-Up Plan
        doc.add_heading("6. Follow-Up Plan", level=1)
        followup = signal_data.get("followup_plan", {})
        
        steps = followup.get("steps", [])
        if isinstance(steps, list):
            for step in steps:
                doc.add_paragraph(f"• {step}", style='List Bullet')
        else:
            doc.add_paragraph(str(steps))
        
        priority = followup.get("priority", "Medium")
        timeline = followup.get("timeline", "")
        
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("Priority: ").bold = True
        p.add_run(priority)
        
        if timeline:
            p = doc.add_paragraph()
            p.add_run("Timeline: ").bold = True
            p.add_run(timeline)
        
        doc.add_paragraph("")
        
        # Section 7: AI Summary
        doc.add_heading("7. AI Summary (Regulatory Style)", level=1)
        ai_summary = signal_data.get("ai_summary", "No summary available.")
        doc.add_paragraph(ai_summary)
        doc.add_paragraph("")
        
        # Save the file
        doc.save(output_path)
        return output_path

